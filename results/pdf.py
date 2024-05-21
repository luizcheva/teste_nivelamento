import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from pathlib import Path
from docx2pdf import convert
from datetime import datetime


class createPDF:
    def __init__(
            self, nome: str, matricula: int, setor: str, list_results: list
    ) -> None:
        self.nome = nome
        self.matricula = matricula
        self.setor = setor
        self.list_results = list_results
        self.setupPDF()

    def changeFillColor(self, cell, color):
        cell_xml = cell._tc
        table_cells_property = cell_xml.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), color)
        table_cells_property.append(shade_obj)

    def resize_image(self, image_path, new_height: float | int = 264.5):
        ROOT_DIR = Path(__file__).parent
        DIR_IMG = ROOT_DIR / 'temp_imagens'
        nome_file, extensao = os.path.splitext(os.path.basename(image_path))
        img = Image.open(image_path)
        width, height = img.size
        new_height = round(new_height)
        new_width = round(width * new_height / height)

        img = img.resize(size=(new_width, new_height))

        # Salve a imagem redimensionada temporariamente
        if not os.path.exists(DIR_IMG):
            os.makedirs(DIR_IMG)

        NEW_IMAGE = DIR_IMG / f'{nome_file}_temp{extensao}'
        img.save(
            NEW_IMAGE,
        )
        return NEW_IMAGE

    def setupPDF(self):
        acerto = 0
        erro = 0
        for status in self.list_results:
            sts = status['STATUS']
            if sts:
                acerto += 1
            else:
                erro += 1
        doc = Document()
        section_1 = doc.sections[0]
        section_1.page_width = Inches(8.27)  # Largura da folha A4
        section_1.page_height = Inches(11.69)  # Altura da folha A4
        section_1.left_margin = Inches(0.2)  # Margem esquerda
        section_1.right_margin = Inches(0.2)  # Margem direita
        section_1.top_margin = Inches(0.2)  # Margem superior
        section_1.bottom_margin = Inches(0.2)  # Margem inferior

        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'

        table.rows[0].height = Inches(0.60)
        table.rows[1].height = Inches(0.34)
        table.rows[2].height = Inches(0.34)

        cell_00 = table.cell(0, 0)
        cell_01 = table.cell(0, 1)
        cell_02 = table.cell(0, 2)

        for cel in table._cells:
            self.changeFillColor(cel, "F3F3F3")

        cell_00.width = Inches(3.94)

        run_1 = cell_00.paragraphs[0].add_run()
        run_1.add_picture('img/neodent-logo.png', Inches(2))

        cell_00.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_00.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER

        cell_01.merge(cell_02)
        cell_01.paragraphs[0].text = 'TESTE DE NIVELAMENTO'
        cell_01.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_01.paragraphs[0].style.font.name = "72 Light"
        cell_01.paragraphs[0].style.font.size = Pt(14)
        cell_01.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER

        cell_10 = table.cell(1, 0)
        cell_11 = table.cell(1, 1)
        cell_12 = table.cell(1, 2)

        cell_10.paragraphs[0].text = f'Colaborador: {self.nome}'
        cell_11.paragraphs[0].text = f'Matrícula: {self.matricula}'
        cell_12.paragraphs[0].text = f'Setor: {self.setor}'

        for cel_edit in table.row_cells(1):
            cel_edit.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            cel_edit.paragraphs[0].style.font.name = "72 Light"
            cel_edit.paragraphs[0].style.font.size = Pt(11)

        cell_20 = table.cell(2, 0)
        cell_21 = table.cell(2, 1)
        cell_22 = table.cell(2, 2)

        cell_21.merge(cell_22)
        cell_20.paragraphs[0].text = f'Total de acertos: {acerto}'
        cell_21.paragraphs[0].text = f'Total de erros: {erro}'
        green = RGBColor(0, 128, 0)
        red = RGBColor(255, 0, 0)
        table.rows[2].cells[0].paragraphs[0].runs[0].font.color.rgb = green
        table.rows[2].cells[1].paragraphs[0].runs[0].font.color.rgb = red

        for cel_edit in table.row_cells(2):
            cel_edit.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
            cel_edit.paragraphs[0].style.font.name = "72 Light"
            cel_edit.paragraphs[0].style.font.size = Pt(11)

        self.changeFillColor(cell_20, '#E2EFDA')
        self.changeFillColor(cell_21, '#FFC5C5')

        doc.add_paragraph()

        grafico1 = doc.add_paragraph()
        grafico1.add_run().add_picture(
            f'results/graficos/{self.nome}_{self.matricula}/{self.setor}/'
            'barGraph.png',
            width=Inches(7)
        )
        grafico1.alignment = WD_TABLE_ALIGNMENT.CENTER

        grafico2 = doc.add_paragraph()
        grafico2.add_run().add_picture(
            f'results/graficos/{self.nome}_{self.matricula}/{self.setor}/'
            'barHGraph.png',
            width=Inches(7)
        )
        grafico2.alignment = WD_TABLE_ALIGNMENT.CENTER

        grafico3 = doc.add_paragraph()
        grafico3.add_run().add_picture(
            f'results/graficos/{self.nome}_{self.matricula}/{self.setor}/'
            'pieGraph.png',
            width=Inches(7)
        )
        grafico3.alignment = WD_TABLE_ALIGNMENT.CENTER

        for num, questions in enumerate(self.list_results):
            table_pergunta = doc.add_table(rows=4, cols=1)
            tbl = table_pergunta._tbl
            for i, tbCel in enumerate(tbl.iter_tcs()):
                if i == 0:
                    tcpr = tbCel.tcPr
                    tcBorders = OxmlElement('w:tcBorders')
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    top.set(qn('w:sz'), '20')

                    tcBorders.append(top)
                    tcpr.append(tcBorders)

            cell_question = table_pergunta.cell(0, 0)
            cell_imagem = table_pergunta.cell(1, 0)
            cell_options = table_pergunta.cell(2, 0)

            pergunta = questions['PERGUNTA']
            categoria = questions['CATEGORIA_PERGUNTA']

            cell_question.paragraphs[0].text = (
                f'{num + 1}) {pergunta}\n'
                f'(Categoria da Pergunta: {categoria})'
            )

            img_question = questions['IMAGEM_PERGUNTA']
            resposta_correta = questions['RESPOSTA_CORRETA']
            list_images = img_question.split(';')
            if list_images == 1:
                caminho_img = self.resize_image(img_question)
                cell_imagem.paragraphs[0].add_run().add_picture(
                    str(caminho_img),
                    height=Inches(2.79)
                )
                os.remove(caminho_img)
                cell_imagem.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            else:
                table_imgs = cell_imagem.add_table(1, len(list_images))
                for i, imgs in enumerate(list_images):
                    caminho = imgs.strip()
                    caminho_img = self.resize_image(caminho)
                    table_imgs.cell(0, i).paragraphs[0].add_run().add_picture(
                        str(caminho_img),
                        height=Inches(2.79)
                    )
                    os.remove(caminho_img)
                    p = table_imgs.cell(0, i).paragraphs[0]
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER

            list_options = questions['ALTERNATIVAS'].split(';')
            for num, option in enumerate(list_options):
                option = option.strip()
                if option == resposta_correta:
                    if num == 0:
                        cell_options.paragraphs[num].text = f'(X) {option}'
                    else:
                        cell_options.add_paragraph(f'(X) {option}')
                else:
                    if num == 0:
                        cell_options.paragraphs[num].text = f'( ) {option}'
                    else:
                        cell_options.add_paragraph(f'( ) {option}')

            status_question = questions['STATUS']
            runs_status_par = table_pergunta.rows[3].cells[0].paragraphs[0]
            runs_status = runs_status_par.add_run()
            runs_status.font.name = "Calibri"
            runs_status.font.size = Pt(12)
            if status_question:
                runs_status.text = 'Você acertou!'
                runs_status.font.color.rgb = green
            else:
                runs_status.text = 'Você errou!'
                runs_status.font.color.rgb = red

            for cel_edit in table_pergunta._cells:
                cel_edit.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER
                cel_edit.paragraphs[0].style.font.name = "Calibri"
                cel_edit.paragraphs[0].style.font.size = Pt(11)

        agora = datetime.now()
        agora = agora.strftime("%Y-%m-%d %H_%M_%S")
        name_file = (
            f'results/DOC/Teste de Nivelamento - {self.nome}_{self.matricula}_'
            f'{self.setor}_{agora}.docx'
        )
        doc.save(name_file)
        convert(name_file, 'results/PDF')
        return 'Concluido com sucesso!'


if __name__ == '__main__':
    pdf = createPDF()
    print(pdf)
